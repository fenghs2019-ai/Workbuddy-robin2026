#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
将童话诗文档拆分成162个独立文档
"""

import os
import re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn

# 设置页面函数
def set_page_margins(doc, top=1.0, bottom=1.0, left=1.25, right=1.25):
    """设置页边距（单位：英寸）"""
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(top)
        section.bottom_margin = Inches(bottom)
        section.left_margin = Inches(left)
        section.right_margin = Inches(right)

# 设置页脚函数
def add_footer(doc, text):
    """添加页脚"""
    section = doc.sections[0]
    footer = section.footer
    paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    paragraph.text = text
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 设置字体大小
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    run.font.size = Pt(10.5)
    run.font.name = '宋体'

# 解析故事内容
def parse_stories(content_file):
    """解析提取的故事内容"""
    with open(content_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    stories = []
    current_story = None
    
    for i, line in enumerate(lines):
        content = line.strip()
        if ':' in content:
            # 提取冒号后的内容
            text = content.split(':', 1)[1].strip()
            # 检查是否是故事标题（如 "01.羽"）
            if text and text[0].isdigit() and '.' in text and len(text.split('.')[0]) <= 3:
                # 保存前一个故事
                if current_story:
                    stories.append(current_story)
                
                # 开始新故事
                match = re.match(r'(\d+)\.(.+)', text)
                if match:
                    num = match.group(1)
                    title = match.group(2)
                    current_story = {
                        'num': num,
                        'title': title,
                        'full_title': text,
                        'lines': [],
                        'start_line': i + 1
                    }
            else:
                # 是故事内容行
                if current_story:
                    current_story['lines'].append(text)
        elif current_story:
            # 是故事内容行（没有行号的行）
            if content:  # 非空行
                current_story['lines'].append(content)
    
    # 保存最后一个故事
    if current_story:
        stories.append(current_story)
    
    return stories

# 创建单个故事文档
def create_story_doc(story, output_dir):
    """为单个故事创建独立文档"""
    # 创建文档
    doc = Document()
    
    # 设置页边距（A4纸，上下2.54cm约1英寸，左右3.17cm约1.25英寸）
    set_page_margins(doc, top=1.0, bottom=1.0, left=1.25, right=1.25)
    
    # 添加标题
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(story['full_title'])
    title_run.font.name = '微软雅黑'
    title_run.font.size = Pt(16)
    title_run.font.bold = True
    
    # 设置中文字体
    title_run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    # 添加空行
    doc.add_paragraph()
    
    # 添加故事内容
    for line in story['lines']:
        if line.strip():  # 非空行
            content_para = doc.add_paragraph()
            content_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            content_run = content_para.add_run(line)
            content_run.font.name = '宋体'
            content_run.font.size = Pt(12)
            # 设置中文字体
            content_run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        else:
            # 空行
            doc.add_paragraph()
    
    # 添加页脚
    add_footer(doc, '森林的孩子·给小孩子·故事诗')
    
    # 保存文档
    filename = f"{story['num']}-{story['title']}.docx"
    filepath = os.path.join(output_dir, filename)
    doc.save(filepath)
    
    return filepath

# 主函数
def main():
    # 输入文件
    content_file = 'extracted_content.txt'
    
    # 输出目录
    output_dir = '森林的孩子-故事诗拆分版'
    
    # 创建输出目录
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
        print(f"创建输出目录: {output_dir}")
    
    # 解析故事
    print("正在解析故事...")
    stories = parse_stories(content_file)
    print(f"找到 {len(stories)} 个故事")
    
    # 创建文档
    print(f"\n正在创建文档到 {output_dir} ...")
    success_count = 0
    error_count = 0
    
    for i, story in enumerate(stories, 1):
        try:
            filepath = create_story_doc(story, output_dir)
            print(f"[{i}/{len(stories)}] 创建成功: {os.path.basename(filepath)}")
            success_count += 1
        except Exception as e:
            print(f"[{i}/{len(stories)}] 创建失败: {story['full_title']} - {str(e)}")
            error_count += 1
    
    # 打印统计
    print(f"\n{'='*60}")
    print(f"拆解完成！")
    print(f"{'='*60}")
    print(f"总故事数: {len(stories)}")
    print(f"成功创建: {success_count} 个文档")
    print(f"创建失败: {error_count} 个文档")
    print(f"输出目录: {os.path.abspath(output_dir)}")
    print(f"{'='*60}")

if __name__ == '__main__':
    main()
