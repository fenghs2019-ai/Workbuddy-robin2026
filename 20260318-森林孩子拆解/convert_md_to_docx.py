from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import re

def convert_markdown_to_docx(md_file_path, docx_file_path):
    # 读取Markdown文件
    with open(md_file_path, 'r', encoding='utf-8') as f:
        md_content = f.read()

    # 创建Word文档
    doc = Document()

    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Microsoft YaHei'
    font.size = Pt(12)

    # 设置页面边距
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # 按行处理Markdown内容
    lines = md_content.split('\n')
    i = 0

    while i < len(lines):
        line = lines[i].strip()

        if not line:
            i += 1
            continue

        # 处理一级标题 (#)
        if line.startswith('# '):
            title = line[2:].strip()
            heading = doc.add_heading(title, level=1)
            heading.style.font.size = Pt(18)
            heading.style.font.bold = True

        # 处理二级标题 (##)
        elif line.startswith('## '):
            title = line[3:].strip()
            heading = doc.add_heading(title, level=2)
            heading.style.font.size = Pt(16)
            heading.style.font.bold = True

        # 处理三级标题 (###)
        elif line.startswith('### '):
            title = line[4:].strip()
            heading = doc.add_heading(title, level=3)
            heading.style.font.size = Pt(14)
            heading.style.font.bold = True

        # 处理四级标题 (####)
        elif line.startswith('#### '):
            title = line[5:].strip()
            heading = doc.add_heading(title, level=4)
            heading.style.font.size = Pt(12)
            heading.style.font.bold = True

        # 处理无序列表 (- 或 *)
        elif line.startswith('- ') or line.startswith('* '):
            content = line[2:].strip()
            p = doc.add_paragraph(content, style='List Bullet')

        # 处理有序列表 (1. 2. 等)
        elif re.match(r'^\d+\.\s', line):
            content = re.sub(r'^\d+\.\s', '', line).strip()
            p = doc.add_paragraph(content, style='List Number')

        # 处理分隔线 (--- 或 ***)
        elif line == '---' or line == '***':
            p = doc.add_paragraph()
            p.add_run('_' * 80)

        # 处理普通段落
        else:
            # 移除Markdown格式标记
            clean_text = line

            # 移除加粗标记 **text**
            clean_text = re.sub(r'\*\*(.*?)\*\*', r'\1', clean_text)

            # 移除斜体标记 *text*
            clean_text = re.sub(r'\*(.*?)\*', r'\1', clean_text)

            # 移除行内代码标记 `text`
            clean_text = re.sub(r'`(.*?)`', r'\1', clean_text)

            # 移除链接标记 [text](url)
            clean_text = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', clean_text)

            if clean_text:
                p = doc.add_paragraph(clean_text)
                p.paragraph_format.line_spacing = 1.5

        i += 1

    # 保存Word文档
    doc.save(docx_file_path)
    print(f'Word文档已成功保存到: {docx_file_path}')

if __name__ == '__main__':
    md_file = '腾讯AI+游戏战略布局报告.md'
    docx_file = '腾讯AI+游戏战略布局报告.docx'

    try:
        convert_markdown_to_docx(md_file, docx_file)
    except FileNotFoundError:
        print(f'错误: 找不到文件 {md_file}')
    except Exception as e:
        print(f'转换过程中发生错误: {str(e)}')
