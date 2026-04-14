#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
验证生成的文档
"""

import os
from docx import Document

output_dir = '森林的孩子-故事诗拆分版'

# 验证前5个文档
sample_files = ['01-羽.docx', '02-桥.docx', '03-秋.docx', '100-爱.docx', '162-别.docx']

print("="*60)
print("文档验证报告")
print("="*60)

for filename in sample_files:
    filepath = os.path.join(output_dir, filename)
    print(f"\n文件: {filename}")
    print("-"*60)

    try:
        doc = Document(filepath)

        # 检查段落
        print(f"段落数量: {len(doc.paragraphs)}")

        # 显示前几个段落
        for i, para in enumerate(doc.paragraphs[:10]):
            if para.text.strip():
                text = para.text[:50] + '...' if len(para.text) > 50 else para.text
                print(f"  段落{i+1}: {text}")

        # 检查页脚
        if doc.sections and doc.sections[0].footer:
            footer = doc.sections[0].footer.paragraphs[0]
            print(f"页脚: {footer.text if footer else '无'}")

    except Exception as e:
        print(f"错误: {str(e)}")

print("\n" + "="*60)
print(f"文件夹中的文档总数: {len([f for f in os.listdir(output_dir) if f.endswith('.docx')])}")
print("="*60)
